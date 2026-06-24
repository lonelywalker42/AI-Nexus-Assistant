"""导出服务 — DOCX / Markdown 参考文献列表"""

import json
import re
from typing import Optional
from sqlalchemy.orm import Session


def _check_docx_available() -> bool:
    """检查 python-docx 是否可用"""
    try:
        import docx
        return True
    except ImportError:
        return False


def _set_run_font(run, ascii_font="Times New Roman", east_asian_font="宋体", size=None):
    """设置 Run 的字体（同时设置西文和东亚字体）"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    run.font.name = ascii_font
    # 设置东亚字体（中文）
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), east_asian_font)
    if size:
        run.font.size = size


def _add_formatted_runs(paragraph, text, base_bold=False, base_italic=False,
                        ascii_font="Times New Roman", east_asian_font="宋体", size=None):
    """解析行内 Markdown 格式并向 paragraph 添加带格式的 Runs。

    支持: **粗体**, *斜体*, `行内代码`, [链接文本](url)
    """
    from docx.shared import Pt

    # 按优先级拆分行内元素
    # 模式: **bold** | *italic* | `code` | [text](url) | 普通文本
    pattern = re.compile(
        r'\*\*(.+?)\*\*'       # 粗体
        r'|(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)'  # 斜体（避免匹配 **）
        r'|`(.+?)`'            # 行内代码
        r'|\[([^\]]+)\]\(([^)]+)\)'  # 链接
        r'|([^*`\[]+)'         # 普通文本
    )

    for m in pattern.finditer(text):
        bold_text = m.group(1)
        italic_text = m.group(2)
        code_text = m.group(3)
        link_text = m.group(4)
        link_url = m.group(5)
        plain_text = m.group(6)

        if bold_text is not None:
            run = paragraph.add_run(bold_text)
            run.bold = True
            run.italic = base_italic
            _set_run_font(run, ascii_font, east_asian_font, size)
        elif italic_text is not None:
            run = paragraph.add_run(italic_text)
            run.bold = base_bold
            run.italic = True
            _set_run_font(run, ascii_font, east_asian_font, size)
        elif code_text is not None:
            run = paragraph.add_run(code_text)
            run.bold = base_bold
            run.italic = base_italic
            _set_run_font(run, "Courier New", east_asian_font, Pt(10) if size is None else size)
            # 行内代码背景色
            from docx.oxml.ns import qn
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                rPr = run._element.makeelement(qn('w:rPr'), {})
                run._element.insert(0, rPr)
            shd = rPr.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): 'F0F0F0',
            })
            rPr.append(shd)
        elif link_text is not None:
            # 链接显示为 "文本 (url)" 格式
            run = paragraph.add_run(f"{link_text} ({link_url})")
            run.bold = base_bold
            run.italic = base_italic
            run.font.color.rgb = None  # 继承默认颜色
            _set_run_font(run, ascii_font, east_asian_font, size)
            # 设置蓝色下划线
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            run.underline = True
        elif plain_text is not None:
            run = paragraph.add_run(plain_text)
            run.bold = base_bold
            run.italic = base_italic
            _set_run_font(run, ascii_font, east_asian_font, size)


def export_docx(content: str, output_path: str, title: str = "文档") -> dict:
    """将 Markdown 内容导出为 DOCX。

    改进:
      - 中文字体支持（宋体 + Times New Roman）
      - 行内格式解析（粗体、斜体、行内代码、链接）
      - 标题/列表中的行内格式正确渲染
    """
    if not _check_docx_available():
        return {"status": "error", "message": "python-docx 未安装，请运行: pip install python-docx"}

    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 设置默认字体（Normal 样式）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    # 设置东亚字体
    from docx.oxml.ns import qn
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = style.element.makeelement(qn('w:rPr'), {})
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置标题样式的东亚字体
    for lvl in range(1, 4):
        try:
            heading_style = doc.styles[f'Heading {lvl}']
            hRPr = heading_style.element.find(qn('w:rPr'))
            if hRPr is None:
                hRPr = heading_style.element.makeelement(qn('w:rPr'), {})
                heading_style.element.append(hRPr)
            hRFonts = hRPr.find(qn('w:rFonts'))
            if hRFonts is None:
                hRFonts = hRPr.makeelement(qn('w:rFonts'), {})
                hRPr.insert(0, hRFonts)
            hRFonts.set(qn('w:eastAsia'), '宋体')
        except KeyError:
            pass

    # 解析 Markdown 并添加到文档
    lines = content.split('\n')
    in_code_block = False
    code_content = []

    for line in lines:
        # 代码块
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                if code_content:
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run('\n'.join(code_content))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    # 设置东亚字体
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = run._element.makeelement(qn('w:rPr'), {})
                        run._element.insert(0, rPr)
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = rPr.makeelement(qn('w:rFonts'), {})
                        rPr.insert(0, rFonts)
                    rFonts.set(qn('w:eastAsia'), '宋体')
                    # 代码块背景色
                    from docx.oxml import OxmlElement
                    pPr = p._element.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'F5F5F5')
                    pPr.append(shd)
                code_content = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_content.append(line)
            continue

        # 标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if level <= 3:
                heading = doc.add_heading(level=level)
                _add_formatted_runs(heading, text, base_bold=True)
            else:
                p = doc.add_paragraph()
                _add_formatted_runs(p, text, base_bold=True)
            continue

        # 无序列表
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, text)
            continue

        # 有序列表
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            _add_formatted_runs(p, text)
            continue

        # 普通段落（含行内格式）
        if line.strip():
            p = doc.add_paragraph()
            _add_formatted_runs(p, line)
        else:
            doc.add_paragraph()

    try:
        doc.save(output_path)
        return {"status": "ok", "path": output_path}
    except Exception as e:
        return {"status": "error", "message": f"DOCX 保存失败: {str(e)}"}


def export_markdown_refs(db: Session, paper_ids: list[str] = None,
                         style: str = "gb7714") -> str:
    """导出 Markdown 参考文献列表"""
    from app.models.paper import Paper
    from app.search.citation import format_gb

    if paper_ids:
        papers = db.query(Paper).filter(Paper.id.in_(paper_ids)).all()
    else:
        papers = db.query(Paper).all()

    if not papers:
        return ""

    refs = []
    for i, p in enumerate(papers, 1):
        paper_dict = {
            "title": p.title,
            "authors": json.loads(p.authors) if p.authors else [],
            "year": p.year,
            "doi": p.doi,
            "journal": p.journal,
            "paper_type": p.paper_type,
        }

        if style == "gb7714":
            ref = format_gb(paper_dict, i)
        elif style == "apa":
            ref = _format_apa(paper_dict)
        elif style == "ieee":
            ref = _format_ieee(paper_dict, i)
        elif style == "mla":
            ref = _format_mla(paper_dict)
        else:
            ref = format_gb(paper_dict, i)

        refs.append(f"[{i}] {ref}")

    return "\n\n".join(refs)


def _format_apa(paper: dict) -> str:
    """APA 7th edition"""
    authors = paper.get("authors", [])
    year = paper.get("year", "n.d.")
    title = paper.get("title", "")
    journal = paper.get("journal", "")
    doi = paper.get("doi", "")

    if not authors:
        author_str = "Unknown Author"
    elif len(authors) == 1:
        parts = authors[0].split()
        author_str = f"{parts[-1]}, {' '.join(p[0] + '.' for p in parts[:-1])}" if len(parts) > 1 else authors[0]
    elif len(authors) <= 20:
        formatted = []
        for a in authors:
            parts = a.split()
            if len(parts) > 1:
                formatted.append(f"{parts[-1]}, {' '.join(p[0] + '.' for p in parts[:-1])}")
            else:
                formatted.append(a)
        author_str = ", ".join(formatted[:-1]) + ", & " + formatted[-1]
    else:
        parts = authors[0].split()
        first = f"{parts[-1]}, {' '.join(p[0] + '.' for p in parts[:-1])}" if len(parts) > 1 else authors[0]
        author_str = f"{first}, ... {authors[-1].split()[-1]}"

    s = f"{author_str} ({year}). {title}."
    if journal:
        s += f" *{journal}*."
    if doi:
        s += f" https://doi.org/{doi}"
    return s


def _format_ieee(paper: dict, idx: int = 1) -> str:
    """IEEE"""
    authors = paper.get("authors", [])
    year = paper.get("year", "")
    title = paper.get("title", "")
    journal = paper.get("journal", "")

    formatted = []
    for a in authors:
        parts = a.split()
        if len(parts) > 1:
            formatted.append(f"{' '.join(p[0] + '.' for p in parts[:-1])} {parts[-1]}")
        else:
            formatted.append(a)
    author_str = ", ".join(formatted) if formatted else "Unknown"

    s = f"[{idx}] {author_str}, \"{title}\""
    if journal:
        s += f", *{journal}*"
    if year:
        s += f", {year}"
    s += "."
    return s


def _format_mla(paper: dict) -> str:
    """MLA 9th edition"""
    authors = paper.get("authors", [])
    year = paper.get("year", "")
    title = paper.get("title", "")
    journal = paper.get("journal", "")

    if not authors:
        author_str = "Unknown Author"
    elif len(authors) == 1:
        parts = authors[0].split()
        author_str = f"{parts[-1]}, {' '.join(parts[:-1])}" if len(parts) > 1 else authors[0]
    elif len(authors) == 2:
        parts = authors[0].split()
        first = f"{parts[-1]}, {' '.join(parts[:-1])}" if len(parts) > 1 else authors[0]
        author_str = f"{first}, and {authors[1]}"
    else:
        parts = authors[0].split()
        author_str = f"{parts[-1]}, {' '.join(parts[:-1])}, et al." if len(parts) > 1 else f"{authors[0]}, et al."

    s = f'{author_str}. "{title}."'
    if journal:
        s += f" *{journal}*"
    if year:
        s += f", {year}"
    s += "."
    return s
